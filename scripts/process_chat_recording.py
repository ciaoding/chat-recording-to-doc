#!/usr/bin/env python3
import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
import time
from difflib import SequenceMatcher
from pathlib import Path
from typing import Optional

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
except ImportError:
    Document = None


TIME_RE = re.compile(r"^(上午|下午|晚上|凌晨|中午)?\d{1,2}:\d{2}$")
VOICE_RE = re.compile(r"(?<!\d)(\d{1,3})\s*[\"”秒](?!\d)")
PHONE_RE = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
ID_RE = re.compile(r"(?<!\d)(?:\d{17}[\dXx]|\d{15})(?!\d)")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
CREDIT_CODE_RE = re.compile(r"(?<![0-9A-Z])[0-9A-Z]{18}(?![0-9A-Z])")
BANK_RE = re.compile(r"(?<!\d)\d{12,19}(?!\d)")

DROP_PATTERNS = [
    re.compile(r"^\d{1,2}:\d{2}"),
    re.compile(r"^[<〈]$"),
    re.compile(r"按住\s*说话"),
    re.compile(r"^[+©g]$"),
    re.compile(r"^[•。]+$"),
    re.compile(r"^PDF$"),
    re.compile(r"KB$"),
    re.compile(r"^m?o?ll", re.I),
]


def stable_token(entity_type: str, value: str, salt: str) -> str:
    digest = hashlib.sha256((salt + "\0" + entity_type + "\0" + value).encode("utf-8")).hexdigest()[:8]
    return f"[{entity_type}_{digest}]"


def normalize(text: str) -> str:
    text = re.sub(r"\s+", "", text or "")
    return text.replace("，", ",").replace("。", ".").replace("：", ":").replace("…", "")


def similarity(a: str, b: str) -> float:
    return SequenceMatcher(None, normalize(a), normalize(b)).ratio()


def run_ocr(video: Path, output: Path, step: float, languages: str, force: bool):
    ocr_jsonl = output / "ocr.jsonl"
    if ocr_jsonl.exists() and not force:
        print(f"Using existing OCR: {ocr_jsonl}")
        return ocr_jsonl
    swift_script = Path(__file__).with_name("ocr_video.swift")
    cmd = [
        "swift",
        str(swift_script),
        "--video",
        str(video),
        "--out",
        str(output),
        "--step",
        str(step),
        "--languages",
        languages,
    ]
    subprocess.run(cmd, check=True)
    return ocr_jsonl


def should_drop(line: dict) -> bool:
    text = line["text"].strip()
    y = line["y"]
    if y < 120 or y > 1190:
        return True
    if TIME_RE.match(text):
        return False
    return any(pattern.search(text) for pattern in DROP_PATTERNS)


def classify_side_by_position(line: dict, width: int) -> str:
    center = line["x"] + line["w"] / 2
    if center > width * 0.56:
        return "right"
    if center < width * 0.44:
        return "left"
    return "center"


def classify_side_by_color(image_path: str, line: dict, width: int) -> str:
    if Image is None:
        return classify_side_by_position(line, width)
    try:
        image = Image.open(image_path).convert("RGB")
    except Exception:
        return classify_side_by_position(line, width)
    x, y, w, h = line["x"], line["y"], line["w"], line["h"]
    x0 = max(0, int(x - 26))
    y0 = max(0, int(y - 16))
    x1 = min(image.width, int(x + w + 26))
    y1 = min(image.height, int(y + h + 16))
    crop = image.crop((x0, y0, x1, y1))
    if crop.width < 1 or crop.height < 1:
        return classify_side_by_position(line, width)
    small = crop.resize((max(1, crop.width // 4), max(1, crop.height // 4)))
    green = white = total = 0
    for r, g, b in small.getdata():
        total += 1
        if g > 175 and r < 195 and b < 180:
            green += 1
        if r > 225 and g > 225 and b > 225:
            white += 1
    if green / max(total, 1) > 0.08:
        return "right"
    if white / max(total, 1) > 0.18:
        return "left"
    return classify_side_by_position(line, width)


def frame_groups(frame: dict) -> list[dict]:
    lines = sorted(frame["lines"], key=lambda line: (line["y"], line["x"]))
    current_time = None
    usable = []
    for line in lines:
        text = line["text"].strip()
        if not text:
            continue
        if should_drop(line):
            continue
        if TIME_RE.match(text):
            current_time = text
            continue

        side = classify_side_by_color(frame["image"], line, frame["width"])
        voice_match = VOICE_RE.search(text)
        is_voice = bool(voice_match and len(normalize(text)) <= 8)
        if is_voice:
            text = f"[语音 {voice_match.group(1)}秒]"

        usable.append(
            {
                "text": text,
                "side": side,
                "chat_time": current_time,
                "frame": frame["frame"],
                "video_time": frame["time"],
                "x": line["x"],
                "y": line["y"],
                "w": line["w"],
                "h": line["h"],
                "is_voice": is_voice,
            }
        )

    groups = []
    for item in usable:
        if not groups:
            groups.append({**item, "texts": [item["text"]], "ys": [item["y"]], "is_voice": item["is_voice"]})
            continue
        prev = groups[-1]
        close = item["y"] - max(prev["ys"]) < 62
        same_side = item["side"] == prev["side"]
        if same_side and close and not prev["is_voice"] and not item["is_voice"]:
            prev["texts"].append(item["text"])
            prev["ys"].append(item["y"])
            prev["text"] = "\n".join(prev["texts"])
            prev["chat_time"] = item.get("chat_time") or prev.get("chat_time")
        else:
            groups.append({**item, "texts": [item["text"]], "ys": [item["y"]], "is_voice": item["is_voice"]})
    return groups


def merge_repeats(groups: list[dict]) -> list[dict]:
    messages = []
    for group in groups:
        text = group["text"].strip()
        if not text or len(normalize(text)) < 2:
            continue
        found = None
        for idx in range(max(0, len(messages) - 24), len(messages)):
            msg = messages[idx]
            if msg["side"] != group["side"]:
                continue
            score = similarity(msg["text"], text)
            if score > 0.84 or normalize(msg["text"]) in normalize(text) or normalize(text) in normalize(msg["text"]):
                found = idx
                break
        if found is None:
            messages.append(group)
        else:
            old = messages[found]
            if len(normalize(text)) > len(normalize(old["text"])):
                group["chat_time"] = group.get("chat_time") or old.get("chat_time")
                messages[found] = group
    return messages


def attach_visible_voice_transcripts(messages: list[dict]) -> list[dict]:
    result = []
    skip = set()
    for i, msg in enumerate(messages):
        if i in skip:
            continue
        if msg.get("is_voice"):
            attached = []
            for j in range(i + 1, min(len(messages), i + 4)):
                candidate = messages[j]
                if candidate.get("is_voice"):
                    continue
                if candidate["side"] != msg["side"]:
                    continue
                if candidate["video_time"] - msg["video_time"] > 3.5:
                    continue
                if len(normalize(candidate["text"])) < 4:
                    continue
                attached.append(candidate["text"])
                skip.add(j)
            if attached:
                msg = dict(msg)
                msg["text"] = msg["text"] + "\n[语音转文字] " + "\n".join(attached)
        result.append(msg)
    return result


def ask_ollama_entities(text: str, model: str, host: str, timeout: int):
    try:
        import requests
    except ImportError:
        return []
    session = requests.Session()
    session.trust_env = False
    prompt = f"""请从以下聊天记录文本中识别需要脱敏的实体。只返回 JSON 数组。
每个元素格式：{{"type":"姓名/公司/地址/项目/账号/微信号/其他标识","value":"原文"}}
不要返回金额、日期、普通语气词。value 必须逐字出现在文本中。

文本：
{text[:12000]}
"""
    try:
        response = session.post(
            f"{host.rstrip('/')}/api/generate",
            json={"model": model, "prompt": prompt, "stream": False},
            timeout=timeout,
        )
        response.raise_for_status()
    except Exception:
        return []
    raw = response.json().get("response", "")
    start = raw.find("[")
    end = raw.rfind("]")
    if start == -1 or end <= start:
        return []
    try:
        value = json.loads(raw[start : end + 1])
    except json.JSONDecodeError:
        return []
    return value if isinstance(value, list) else []


def anonymize_text(text: str, salt: str, model: Optional[str], host: str, timeout: int):
    entities = []
    regexes = [
        ("手机号", PHONE_RE),
        ("身份证号", ID_RE),
        ("邮箱", EMAIL_RE),
        ("统一社会信用代码", CREDIT_CODE_RE),
        ("银行账号", BANK_RE),
    ]
    for entity_type, pattern in regexes:
        for match in pattern.finditer(text):
            entities.append({"type": entity_type, "value": match.group(0), "source": "regex"})
    if model:
        for entity in ask_ollama_entities(text, model, host, timeout):
            if isinstance(entity, dict) and entity.get("value") in text:
                entity["source"] = "ollama"
                entities.append(entity)

    seen = set()
    clean_entities = []
    for entity in entities:
        entity_type = str(entity.get("type", "敏感标识")).strip() or "敏感标识"
        value = str(entity.get("value", "")).strip()
        if len(value) < 2 or (entity_type, value) in seen:
            continue
        seen.add((entity_type, value))
        clean_entities.append({"type": entity_type, "value": value, "source": entity.get("source", "unknown")})

    report = []
    for entity in sorted(clean_entities, key=lambda item: len(item["value"]), reverse=True):
        value = entity["value"]
        count = text.count(value)
        if count == 0:
            continue
        token = stable_token(entity["type"], value, salt)
        text = text.replace(value, token)
        report.append({"type": entity["type"], "token": token, "count": count, "source": entity["source"]})
    return text, report


def sender_label(side: str, left_label: str, right_label: str) -> str:
    if side == "right":
        return right_label
    if side == "left":
        return left_label
    return "系统/时间标记"


def build_markdown(messages, path: Path, title: str, date: str, left_label: str, right_label: str, anonymize: bool, salt: str, model: Optional[str], host: str, timeout: int):
    raw_lines = [f"# {title}", "", f"日期：{date}", "", "| 序号 | 时间 | 发送人 | 内容 |", "|---:|---|---|---|"]
    for idx, msg in enumerate(messages, 1):
        text = msg["text"].replace("\n", "<br>")
        raw_lines.append(f"| {idx} | {msg.get('chat_time') or ''} | {sender_label(msg['side'], left_label, right_label)} | {text} |")
    markdown = "\n".join(raw_lines) + "\n"
    report = []
    if anonymize:
        markdown, report = anonymize_text(markdown, salt, model, host, timeout)
    path.write_text(markdown, encoding="utf-8")
    return markdown, report


def build_docx(markdown: str, path: Path):
    if Document is None:
        return False
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)

    lines = markdown.splitlines()
    title = lines[0].lstrip("# ").strip() if lines else "聊天记录"
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    rows = [line for line in lines if line.startswith("|") and not line.startswith("|---")]
    if rows:
        headers = [cell.strip() for cell in rows[0].strip("|").split("|")]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor(20, 40, 70)
        for row_line in rows[1:]:
            values = [cell.strip().replace("<br>", "\n") for cell in row_line.strip("|").split("|")]
            cells = table.add_row().cells
            for idx, value in enumerate(values[: len(cells)]):
                cells[idx].text = value
    else:
        for line in lines[1:]:
            doc.add_paragraph(line)
    doc.save(path)
    return True


def main():
    parser = argparse.ArgumentParser(description="Convert a local chat screen recording to anonymized text documents.")
    parser.add_argument("--video", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--conversation", default="聊天记录")
    parser.add_argument("--date", default=time.strftime("%Y-%m-%d"))
    parser.add_argument("--step", type=float, default=0.25)
    parser.add_argument("--languages", default="zh-Hans,zh-Hant,en-US")
    parser.add_argument("--left-label", default="对方")
    parser.add_argument("--right-label", default="我")
    parser.add_argument("--no-anonymize", action="store_true")
    parser.add_argument("--salt", default=os.environ.get("CHAT_RECORDING_SALT", "chat-recording-to-doc-change-me"))
    parser.add_argument("--ollama-model", default=None, help="Optional local Ollama model for entity detection, e.g. qwen2.5:7b")
    parser.add_argument("--ollama-host", default=os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434"))
    parser.add_argument("--timeout", type=int, default=300)
    parser.add_argument("--force-ocr", action="store_true")
    args = parser.parse_args()

    video = Path(args.video).expanduser().resolve()
    output = Path(args.output).expanduser().resolve()
    output.mkdir(parents=True, exist_ok=True)
    if not video.exists():
        print(f"Video not found: {video}", file=sys.stderr)
        return 2

    ocr_jsonl = run_ocr(video, output, args.step, args.languages, args.force_ocr)
    frames = [json.loads(line) for line in ocr_jsonl.read_text(encoding="utf-8").splitlines() if line.strip()]
    groups = []
    for frame in frames:
        groups.extend(frame_groups(frame))
    messages = attach_visible_voice_transcripts(merge_repeats(groups))

    raw_json = output / "chat_raw_local_only.json"
    raw_json.write_text(json.dumps(messages, ensure_ascii=False, indent=2), encoding="utf-8")

    md_path = output / "chat_anonymized.md"
    markdown, report = build_markdown(
        messages=messages,
        path=md_path,
        title=args.conversation,
        date=args.date,
        left_label=args.left_label,
        right_label=args.right_label,
        anonymize=not args.no_anonymize,
        salt=args.salt,
        model=args.ollama_model,
        host=args.ollama_host,
        timeout=args.timeout,
    )
    json_path = output / "chat_anonymized.json"
    json_path.write_text(json.dumps({"markdown": markdown}, ensure_ascii=False, indent=2), encoding="utf-8")
    report_path = output / "anonymization_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    docx_path = output / "chat_anonymized.docx"
    docx_ok = build_docx(markdown, docx_path)

    print(f"Messages: {len(messages)}")
    print(f"Markdown: {md_path}")
    print(f"JSON: {json_path}")
    if docx_ok:
        print(f"DOCX: {docx_path}")
    else:
        print("DOCX skipped: install python-docx to enable Word output")
    print(f"Local-only raw merge: {raw_json}")
    print(f"Local-only anonymization report: {report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
